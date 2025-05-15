from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from app.services.order_service import get_all_orders
from app.core.database import get_db
from docx import Document
import io
import os

api_router = APIRouter()

TEMPLATE_PATH = "order_template.docx"

def fill_placeholders(paragraphs, data):
    for para in paragraphs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(value or ""))

@api_router.get("/orders/{order_id}.docx")
async def export_single_order_docx(order_id: int, db: AsyncSession = Depends(get_db)):
    orders = await get_all_orders(db)
    
    # 找到對應訂單
    target_order = next((order for order in orders if order.id == order_id), None)
    if not target_order:
        return {"error": "Order not found"}

    # 英文星期對應中文
    WEEKDAY_MAP = {
        "Monday": "星期一",
        "Tuesday": "星期二",
        "Wednesday": "星期三",
        "Thursday": "星期四",
        "Friday": "星期五",
        "Saturday": "星期六",
        "Sunday": "星期日",
    }

    # 對應欄位 mapping（請根據實際資料結構調整）
    mapped_data = {
        # "id": target_order.id,
        "訂購人 | customer_name": target_order.customer_name,
        "訂購人連絡電話 | phone": target_order.customer_phone,
        "timestamp": target_order.order_date.strftime("%Y-%m-%d") if target_order.order_date else "",
        "收據寄送地址": getattr(target_order, "receipt_address", ""),
        "訂購品項 | ITEM": target_order.item,
        "訂購數量 | QTY": target_order.quantity,
        "付款方式 | PAY WAY": getattr(target_order, "payway", ""),
        "備註": target_order.note,
        "卡片內容：（範例請看圖）": getattr(target_order, "card_message", ""),
        "星期": WEEKDAY_MAP.get(
                target_order.send_datetime.strftime("%A"), ""
            ) if getattr(target_order, "send_datetime", None) else "",
        "送禮日期 & 時間 | SEND DAY & TIME": (
            target_order.send_datetime.strftime("%Y-%m-%d %H:%M") if getattr(target_order, "send_datetime", None) else ""
        ),
        "收禮人 / 取貨人全名 | RECEIVER NAME": getattr(target_order, "receiver_name", ""),
        "收禮人連絡電話 | MOBILE": getattr(target_order, "receiver_phone", ""),
        "送禮地址 | ADD": getattr(target_order, "delivery_address", "")
    }

    # 載入模板，並填入欄位（處理表格和段落）
    doc = Document(TEMPLATE_PATH)
    fill_placeholders(doc.paragraphs, mapped_data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fill_placeholders(cell.paragraphs, mapped_data)

    # 輸出
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f"order_{order_id}.docx"
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
